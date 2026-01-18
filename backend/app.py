import os
import whisper
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import google.generativeai as genai
from docx import Document 
import io

# הגדרות
GEMINI_API_KEY = "AIzaSyDzCuawPQUpq9sVaNKCjzVjx0w26dqFtLQ"
genai.configure(api_key=GEMINI_API_KEY)

app = Flask(__name__)
CORS(app)

# טעינת מודל Whisper לזיכרון
print("Loading Whisper Model... please wait.")
model = whisper.load_model("base")
print("Whisper Model Loaded!")

@app.route('/process-meeting', methods=['POST'])
def process_meeting():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files['file']
    # שמירת הקובץ בשם זמני
    file_path = "temp_audio.mp3"
    file.save(file_path)

    try:
        # 1. שלב התמלול (Whisper)
        print("Starting transcription...")
        result = model.transcribe(file_path)
        transcript_text = result['text']

        # 2. שלב הסיכום (Gemini 1.5 Flash - המודל המעודכן)
        print("Starting analysis...")
        gemini_model = genai.GenerativeModel('gemini-flash-latest')
        
        system_instructions = (
            "אתה עוזר ניהולי חכם. נתח את תמלול הפגישה הבא בעברית והוצא דוח מובנה הכולל:\n"
            "1. כותרת לפגישה.\n"
            "2. סיכום תמציתי של עיקרי הדברים.\n"
            "3. רשימת משתתפים (אם הוזכרו).\n"
            "4. החלטות שהתקבלו.\n"
            "5. רשימת משימות לביצוע (Action Items).\n"
            "השתמש בפורמט Markdown ברור."
        )
        
        response = gemini_model.generate_content(f"{system_instructions}\n\nהתמלול:\n{transcript_text}")
        summary_text = response.text

        return jsonify({
            "transcript": transcript_text,
            "summary": summary_text
        })

    except Exception as e:
        print(f"Error during processing: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        # מחיקת הקובץ הזמני כדי לשמור על סדר
        if os.path.exists(file_path):
            os.remove(file_path)

# פונקציה להורדת קובץ Word
@app.route('/download-word', methods=['POST'])
def download_word():
    try:
        data = request.json
        summary = data.get('summary', '')
        transcript = data.get('transcript', '')

        doc = Document()
        # הגדרת כיוון טקסט לימין (חשוב לעברית)
        doc.add_heading('סיכום פגישה אוטומטי', 0)
        
        doc.add_heading('סיכום וניתוח AI', level=1)
        doc.add_paragraph(summary)
        
        doc.add_page_break()
        
        doc.add_heading('תמלול הפגישה', level=1)
        doc.add_paragraph(transcript)

        # שמירת הקובץ לאובייקט בזיכרון (במקום לדיסק)
        target = io.BytesIO()
        doc.save(target)
        target.seek(0)
        
        return send_file(
            target, 
            as_attachment=True, 
            download_name="Meeting_Summary.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"Error generating Word file: {e}")
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    print("--- Server is starting on http://127.0.0.1:5000 ---")
    app.run(port=5000, debug=False)